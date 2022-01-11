from docx import Document
import datetime

def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table


def main(records_morning, records_evening, records_urgent):
    morning, evening, urgent = (), (), ()
    for i in records_morning:
        for j in i:
            morning = (*morning, (j[3], j[4], str(j[5]).replace('1', "+").replace('0', '-')))
    for i in records_evening:
        for j in i:
            evening = (*evening, (j[3], j[4], str(j[5]).replace('1', "+").replace('0', '-')))
    for i in records_urgent:
        for j in i:
            urgent = (*urgent, (j[3], j[4], str(j[5]).replace('1', "+").replace('0', '-')))
    document = Document()
    document.add_heading('Отчет по утренним задачам', 0)
    headers = ('Задача', 'Место', 'Статус выполнения')
    create_table(document, headers, morning)
    document.add_paragraph()
    document.add_heading('Отчет по вечерним задачам', 0)
    create_table(document, headers, evening)
    document.add_paragraph()
    document.add_heading('Отчет по срочным задачам', 0)
    create_table(document, headers, urgent)
    document_name = f"reports/report_{datetime.datetime.now().strftime('%d.%m.%Y_[%H:%M]')}.docx"
    document.save(document_name)
    return document_name
